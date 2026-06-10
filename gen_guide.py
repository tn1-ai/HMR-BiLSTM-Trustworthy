# -*- coding: utf-8 -*-
"""Generate HMR-BiLSTM_Installation_Guide.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

BLUE   = RGBColor(0x1F, 0x49, 0x7D)
RED    = RGBColor(0xC0, 0x00, 0x00)
GREEN  = RGBColor(0x00, 0x60, 0x00)
DKBLUE = RGBColor(0x20, 0x20, 0x80)
GRAY   = RGBColor(0x88, 0x88, 0x88)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(14)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(12)
    return p

def para(text="", bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = GREEN
    return p

def monoline(text):
    """Monospace line for directory trees."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = DKBLUE
    return p

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "1F497D")
    # Rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if r % 2 == 1:
                shade_cell(cell, "DCE6F1")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def spacer():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_heading("INSTALLATION AND SETUP GUIDE", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    run.font.color.rgb = BLUE
    run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Project: HMR-BiLSTM \u2013 A Trustworthy and Explainable "
    "Hybrid Memory Residual BiLSTM Framework for ECG Arrhythmia Classification"
)
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = BLUE
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 1. Overview
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Overview")
para(
    "This guide explains how to set up and run the HMR-BiLSTM project on a "
    "local machine (Windows or Linux/macOS) using Python and PyTorch (CPU or GPU)."
)
para(
    "The project implements a deep learning framework for ECG arrhythmia "
    "classification on the MIT-BIH Arrhythmia Dataset. It includes training, "
    "ablation studies, adversarial robustness evaluation (FGSM & PGD), "
    "calibration analysis, and publication-ready figure/table generation."
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 2. System Requirements
# ══════════════════════════════════════════════════════════════════════════════
h1("2. System Requirements")
add_table(
    ["Component", "Recommended Configuration"],
    [
        ["Platform",       "Windows 10/11 (64-bit) or Linux/macOS"],
        ["GPU (optional)", "NVIDIA GPU with CUDA 12.1+ (e.g. RTX 3060 or higher)"],
        ["Python Version", "3.9 or higher (3.11 recommended)"],
        ["PyTorch",        "2.0.0+"],
        ["RAM",            "8 GB minimum (16 GB recommended)"],
        ["Disk Space",     "~3 GB for dataset, checkpoints, and results"],
        ["CUDA Toolkit",   "12.1 (only if using GPU)"],
    ],
    col_widths=[5.5, 10],
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 3. Required Libraries
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Required Libraries")
para("Install all dependencies using the provided requirements file:")
code("pip install -r requirements.txt")
spacer()
add_table(
    ["Category", "Libraries"],
    [
        ["Deep Learning",    "torch (>=2.0.0), torchvision, torchaudio"],
        ["Data Handling",    "numpy (>=1.24.0), pandas (>=2.0.0)"],
        ["Machine Learning", "scikit-learn (>=1.3.0), imbalanced-learn (>=0.11.0)"],
        ["Visualization",    "matplotlib (>=3.7.0)"],
        ["Built-in (no install)", "os, json, argparse, pathlib"],
    ],
    col_widths=[5.5, 10],
)
spacer()
para("For GPU support (NVIDIA CUDA 12.1), also run:", italic=True, size=10)
code("pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu121")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 4. Project Directory Structure
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Project Directory Structure")
tree = [
    "HMR-BiLSTM/",
    "\u251c\u2500\u2500 data/",
    "\u2502   \u251c\u2500\u2500 raw/                   # Raw MIT-BIH CSV files (git-ignored)",
    "\u2502   \u2514\u2500\u2500 processed/             # Pre-processed splits + class weights",
    "\u251c\u2500\u2500 results/",
    "\u2502   \u251c\u2500\u2500 checkpoints/           # Model checkpoints (.pt)",
    "\u2502   \u251c\u2500\u2500 figures/               # Generated plots and visualizations",
    "\u2502   \u251c\u2500\u2500 tables/                # LaTeX and CSV result tables",
    "\u2502   \u2514\u2500\u2500 logs/                  # Evaluation logs (JSON)",
    "\u251c\u2500\u2500 hmr_bilstm.py              # Model architecture + RLSTMLoss",
    "\u251c\u2500\u2500 hmr_bilstm_ablation.py     # Ablation variants",
    "\u251c\u2500\u2500 preprocess.py              # Data preprocessing and split",
    "\u251c\u2500\u2500 train.py                   # Main training script",
    "\u251c\u2500\u2500 run_baselines.py           # Train baseline models",
    "\u251c\u2500\u2500 run_ablation.py            # Ablation study driver",
    "\u251c\u2500\u2500 report_results.py          # Core figures (confusion matrix, ROC)",
    "\u251c\u2500\u2500 evaluate_fgsm.py           # FGSM robustness evaluation",
    "\u251c\u2500\u2500 evaluate_pgd.py            # PGD robustness evaluation",
    "\u251c\u2500\u2500 evaluate_calibration.py    # Calibration analysis",
    "\u251c\u2500\u2500 evaluate_robustness_all.py # Gaussian noise robustness",
    "\u251c\u2500\u2500 evaluate_ablation_robustness.py",
    "\u251c\u2500\u2500 combine_ablation_tables.py",
    "\u251c\u2500\u2500 generate_results_tables.py",
    "\u251c\u2500\u2500 plot_and_export.py         # Export final figures and tables",
    "\u251c\u2500\u2500 requirements.txt",
    "\u2514\u2500\u2500 README.md",
]
for line in tree:
    monoline(line)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 5. Dataset Access and Setup
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Dataset Access and Setup")
h2("Dataset Overview")
add_table(
    ["Property", "Details"],
    [
        ["Dataset Name",  "MIT-BIH Arrhythmia Dataset (CSV version)"],
        ["Source",        "Kaggle \u2013 shayanfazeli/heartbeat"],
        ["Training size", "109,446 segmented heartbeat samples"],
        ["Test size",     "21,892 segmented heartbeat samples"],
        ["Classes",       "N (Normal), S (Supraventricular), V (Ventricular), F (Fusion), Q (Unknown)"],
        ["Format",        "CSV \u2013 each row is one heartbeat segment (187 features + 1 label column)"],
        ["License",       "Open \u2013 freely available for research"],
        ["Dataset URL",   "https://www.kaggle.com/datasets/shayanfazeli/heartbeat"],
    ],
    col_widths=[5, 10.5],
)
spacer()

h2("Option A: Download via Browser (Recommended)")
steps_a = [
    "Sign in to your Kaggle account.",
    "Visit: https://www.kaggle.com/datasets/shayanfazeli/heartbeat",
    "Click Download to save the .zip file locally.",
    "Extract the ZIP \u2013 you should find mitbih_train.csv and mitbih_test.csv.",
    "Place both CSV files in the data/raw/ folder of this project.",
]
for s in steps_a:
    bullet(s)

spacer()
h2("Option B: Download via Kaggle API (advanced users)")
code("pip install kaggle")
code("kaggle datasets download -d shayanfazeli/heartbeat")
code("unzip heartbeat.zip -d data/raw/")

spacer()
h2("Required Folder Structure After Download")
for line in [
    "HMR-BiLSTM/",
    "\u2514\u2500\u2500 data/",
    "    \u2514\u2500\u2500 raw/",
    "        \u251c\u2500\u2500 mitbih_train.csv",
    "        \u2514\u2500\u2500 mitbih_test.csv",
]:
    monoline(line)
spacer()
para(
    "Note: The data/raw/ folder is git-ignored. Do NOT rename the CSV files.",
    italic=True, size=10, color=RED,
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 6. Setup Steps
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Setup Steps")

h2("Step 1: Clone the Repository")
code("git clone https://github.com/tn1-ai/HMR-BiLSTM.git")
code("cd HMR-BiLSTM")

h2("Step 2: Create a Virtual Environment (Recommended)")
para("Windows:", bold=True, size=10)
code("python -m venv venv")
code(r"venv\Scripts\activate")
para("Linux / macOS:", bold=True, size=10)
code("python -m venv venv")
code("source venv/bin/activate")

h2("Step 3: Install Python Dependencies")
code("pip install -r requirements.txt")
para("Optional \u2013 GPU acceleration (CUDA 12.1):", italic=True, size=10)
code("pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu121")

h2("Step 4: Prepare the Dataset")
para("Place the CSV files in data/raw/ (see Section 5), then run:")
code("python preprocess.py")
para(
    "This creates train/val/test splits and class-weight files under data/processed/.",
    size=11,
)

h2("Step 5: Train the Proposed HMR-BiLSTM Model")
code("python train.py")
para(
    "Trains with adversarial FGSM, focal loss, cosine-annealing LR, and early stopping. "
    "Best checkpoint saved to: results/checkpoints/best_rlstm.pt",
    size=11,
)

h2("Step 6: Train Baseline Models")
code("python run_baselines.py")
para("Trains plain LSTM and BiLSTM baselines. Checkpoints stored in results/checkpoints/.", size=11)

h2("Step 7: Run Ablation Study")
para("Run all variants:", size=11)
code("python run_ablation.py")
para("Run specific variants:", size=11)
code("python run_ablation.py --variants no_rmc no_cnn")
para("Generate table only from existing checkpoints:", size=11)
code("python run_ablation.py --table-only")
para("Available variants: full, no_rmc, no_cnn, mean_pool, no_adv, no_hybrid, no_smooth", italic=True, size=10)

h2("Step 8: Evaluate and Visualize")
add_table(
    ["Command", "Purpose"],
    [
        ["python report_results.py",             "Confusion matrix, ROC curve, gate trajectories"],
        ["python compare_fgsm_baselines.py",      "FGSM adversarial robustness (all models)"],
        ["python evaluate_pgd.py",                "PGD adversarial robustness"],
        ["python evaluate_robustness_all.py",      "Gaussian-noise robustness"],
        ["python evaluate_calibration.py",         "Reliability diagram, ECE, Brier score"],
        ["python evaluate_ablation_robustness.py", "Robustness of each ablation variant"],
        ["python plot_and_export.py",              "Export final figures and LaTeX/CSV tables"],
    ],
    col_widths=[7, 8.5],
)
spacer()
para(
    "Note (Windows/CUDA): evaluate_pgd.py automatically sets "
    "torch.backends.cudnn.enabled = False to prevent cuDNN backward errors.",
    italic=True, size=10, color=RED,
)

h2("Step 9: Export Final Tables and Figures for Paper")
code("python plot_and_export.py")
para("Output files:", size=11)
bullet("Figures: results/figures/ (PDF and PNG)")
bullet("Tables:  results/tables/ (LaTeX .tex and .csv)")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 7. Save and Load Models
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Save and Load Models")
para(
    "After training, all model checkpoints are saved automatically in results/checkpoints/. "
    "To reload a trained model for evaluation:"
)
code("import torch")
code("from hmr_bilstm import HMRBiLSTM")
code("model = HMRBiLSTM(...)  # same hyperparameters as training")
code('checkpoint = torch.load("results/checkpoints/best_rlstm.pt")')
code('model.load_state_dict(checkpoint["model_state_dict"])')
code("model.eval()")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 8. Troubleshooting
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Troubleshooting")
add_table(
    ["Issue", "Possible Fix"],
    [
        ["ModuleNotFoundError",              "Re-run: pip install -r requirements.txt"],
        ["CUDA not detected",                "Verify CUDA toolkit matches PyTorch wheel; reinstall with correct --index-url"],
        ["OutOfMemoryError during training", "Reduce batch size in train.py, or switch to CPU"],
        ["cuDNN backward error (PGD)",       "torch.backends.cudnn.enabled=False is already set in evaluate_pgd.py"],
        ["File not found: mitbih_train.csv", "Ensure CSV files are placed in data/raw/ with exact filenames"],
        ["Low F1 on minority classes",       "Verify class_weights.json exists in data/processed/; re-run preprocess.py"],
        ["Training too slow",                "Use a GPU; or reduce epochs / batch size for quick testing"],
    ],
    col_widths=[6, 9.5],
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 9. Notes
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Notes")
for note in [
    "GPU acceleration is strongly recommended. Training time is approximately 30\u201360 minutes per model variant on a mid-range GPU (e.g., RTX 3060), vs. several hours on CPU.",
    "Ensure your disk has at least 3 GB of free space for dataset, checkpoints, and output results.",
    "All results (checkpoints, logs, figures, tables) are automatically saved under the results/ directory.",
    "The data/raw/ folder is excluded from version control (.gitignore). Never commit the raw dataset.",
    "Re-run preprocess.py any time you replace or update the raw CSV files.",
]:
    bullet(note)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 10. Author
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Author")
para(
    "NGUYEN VAN THONG \u2013 Class FE0001 \u2013 Course 50 \u2013 "
    "UEH College of Technology and Design \u2013 "
    "University of Economics Ho Chi Minh City (UEH)",
    size=11,
)
spacer()
para("Contact:", bold=True, size=11)
bullet("EduMail:  thongnguyen.31241027212@st.ueh.edu.vn")
bullet("Personal: nguyenthong111206@gmail.com")
bullet("GitHub:   https://github.com/tn1-ai")

spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u2500\u2500\u2500 End of Guide \u2500\u2500\u2500")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = GRAY

# ── Save
OUT = r"d:\HMR-BiLSTM-main\HMR-BiLSTM_Installation_Guide.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
